"""DOCX parsing."""

from pathlib import Path


def load_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.lower().startswith("heading"):
            level = style.split()[-1] if style.split()[-1].isdigit() else "1"
            parts.append(f"{'#' * int(level)} {text}")
        else:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)
